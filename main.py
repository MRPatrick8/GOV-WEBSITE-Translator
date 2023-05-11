# import streamlit as st 
# from translator import *
# import pandas as pd
# from accuracy import *
# from docx import Document
# # from annotated_text import annotated_text


# st.title("GOV Contents Translation")

# col1, col2 = st.columns(2)

# with col1:
#    st.image("https://www.developmentaid.org/files/organizationLogos/huzalabs-435178.jpg")

# with col2:
#    st.image("images/Rwanda Gov.png")

# st.image("images/GIZ_HUZALABS_FAIRFORWARD.png")

# # with st.sidebar:
# #     st.header("Choose an option")
# # Add CSS styling to center the element
# st.markdown(
#     """
#     <style>
#     .center {
#         display: flex;
#         justify-content: center;
#         align-items: center;
#         height: 80px;
#         margin: 0 auto;
#     }
#     </style>
#     """,
#     unsafe_allow_html=True,
# )

# # Add the file uploader element inside a div with the center class
# with st.container():
#     st.header("Upload your Docx file")
#     file_upload = st.file_uploader('', type=['docx'])
    

#     #save file file
#     if file_upload:
#         with open('data/content.docx', "wb+") as doc_file:
#             doc_file.write(file_upload.read())
#         #Read the doc file
#         doc = Document('data/content.docx')
#         paragraphs = []
#         for i in range(len(doc.paragraphs)):
#             paragraphs.append(doc.paragraphs[i].text)
#         st.session_state['text_area'] = "\n".join(paragraphs)

#     st.header("Or Paste/Type your Text Here")
#     #Paste in your text docs
#     st.text_area("", key='text_area', height=10)


    
    
# translate_button, download_button, a ,b  = st.columns(4)

# with translate_button:
#     translate_button = st.button("Translate Dataset")
   
       
# if translate_button:
#     paragraphs = st.session_state['text_area'].split('\n')
#     texts = []
#     for text in paragraphs:
#         trans_text = translate_row(text, target_lang='rw', source_lang='en')
#         texts.append(trans_text)

#     #print(texts)
#     st.text_area("Translation", value='\n'.join(texts), disabled=False, height=20)


# with download_button:
#     download_button = st.button("Download Dataset")
import streamlit as st 
from translator import *
import pandas as pd
from accuracy import *
from docx import Document
# from annotated_text import annotated_text


st.title("GOV Contents Translation")

col1, col2 = st.columns(2)

with col1:
   st.image("https://www.developmentaid.org/files/organizationLogos/huzalabs-435178.jpg")

with col2:
   st.image("images/Rwanda Gov.png")

st.image("images/GIZ_HUZALABS_FAIRFORWARD.png")

# with st.sidebar:
#     st.header("Choose an option")
# Add CSS styling to center the element
st.markdown(
    """
    <style>
    .center {
        display: flex;
        justify-content: center;
        align-items: center;
        height: 80px;
        margin: 0 auto;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Add the file uploader element inside a div with the center class
with st.container():
    st.header("Upload your Docx file")
    file_upload = st.file_uploader('', type=['docx'])
    

    #save file file
    if file_upload:
        with open('data/content.docx', "wb+") as doc_file:
            doc_file.write(file_upload.read())
        #Read the doc file
        doc = Document('data/content.docx')
        paragraphs = []
        for i in range(len(doc.paragraphs)):
            paragraphs.append(doc.paragraphs[i].text)
        st.session_state['text_area'] = "\n".join(paragraphs)

    st.header("Or Paste/Type your Text Here")
    #Paste in your text docs
    st.text_area("", key='text_area', height=10)


    
    
from io import BytesIO

translate_button, download_button, a, b = st.columns(4)

with translate_button:
    translate_button = st.button("Translate Dataset")

if translate_button:
    if 'text_area' in st.session_state:
        paragraphs = st.session_state['text_area'].split('\n')
        texts = []
        for text in paragraphs:
            trans_text = translate_row(text, target_lang='rw', source_lang='en')
            texts.append(trans_text)
        if len(texts) > 0:
            st.text_area("Translation", value='\n'.join(texts), disabled=False, height=20)
            st.session_state['translated_texts'] = texts
    else:
        st.warning("Please upload a file or enter text to translate.")

with download_button:
    download_button = st.button("Download Dataset")

if download_button:
    if 'translated_texts' in st.session_state:
        doc = Document()
        for para in st.session_state['translated_texts']:
            doc.add_paragraph(para)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button(
            label="Download Translated Docx File",
            data=doc_io.getvalue(),
            file_name="translated_doc.docx",
            mime="application/octet-stream"
        )
    else:
        st.warning("Please translate a text before downloading the file.")
