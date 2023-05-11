import requests
import json
import string
import pandas as pd
import docx
import re

class GoogleTranslate:
    api_url = 'https://translate.googleapis.com/translate_a/single'
    client = '?client=gtx'
    dt = '&dt=t'

def translate_row(text, target_lang, source_lang):
    sl = f"&sl={source_lang}"
    tl = f"&tl={target_lang}"
    
    # Remove punctuation marks from the input text
    # translator = str.maketrans('', '', string.punctuation)
    # text = row['content'].translate(translator)
    
    r = requests.get(f"{GoogleTranslate.api_url}{GoogleTranslate.client}{GoogleTranslate.dt}{sl}{tl}&q={text}")
    if r.status_code == 200:
        response_data = json.loads(r.text)
        translated_text = response_data[0][0][0]
    else:
        translated_text = ""
    return translated_text

def translate_document(input_file, output_file, target_lang, source_lang):
    # Create a new Word document
    document = docx.Document()
    
    # Read the input document and create a list of paragraphs
    paragraphs = []
    input_doc = docx.Document(input_file)
    for para in input_doc.paragraphs:
        paragraphs.append(para.text)
    
    # Create a dataframe from the paragraphs
    df = pd.DataFrame({'content': paragraphs})
    
    # Translate the dataframe and add the translated text to the dataframe
    df['translated_text'] = translate_dataframe(df, target_lang, source_lang)
    
    # Add the translated text to the output document
    for index, row in df.iterrows():
        document.add_paragraph(row['translated_text'])
    
    # Save the output document
    document.save(output_file)

def translate_paragraphs(input_paragraphs, target_lang, source_lang=None):
    paragraphs = []
    for paragraph in input_paragraphs:
        trans_paragraph = translate_row(paragraph, target_lang, source_lang)
        paragraphs.append(paragraphs)

    return paragraphs

def translate_dataframe(df, target_lang, source_lang):
    df_translated = df.apply(translate_row, axis=1, args=(target_lang, source_lang))
    return df_translated['translated_text']

############################Back##############################################
def back_translate_document(input_file, output_file, target_lang, source_lang):
    # Create a new Word document
    document = docx.Document()
    list_of_contents_modified = [re.sub(r'\s+', ' ', re.sub(r'\W+', ' ', item)).strip() for item in input_file]
    # Read the input document and create a list of paragraphs
    paragraphs = []
    input_doc = docx.Document(input_file)
    for para in input_doc.paragraphs:
        paragraphs.append(para.text)
    
    # Create a dataframe from the paragraphs
    df = pd.DataFrame({'content': paragraphs})
    
    # Translate the dataframe and add the translated text to the dataframe
    df['back_translated_text'] = translate_dataframe(df, target_lang, source_lang)
    
    # Add the translated text to the output document
    for index, row in df.iterrows():
        document.add_paragraph(row['back_translated_text'])
    
    # Save the output document
    document.save(output_file)

def back_translate_dataframe(df, target_lang, source_lang):
    df_translated = df.apply(translate_row, axis=1, args=(target_lang, source_lang))
    return df_translated['back_translated_text']
# Translate the input document and save the output document
# translate_document('content_cleaned.docx', 'output_file.docx', 'rw', 'en')
# back_translate_document('output_file.docx', 'back_output_file.docx', 'en', 'rw')